#!/usr/bin/env python
# -*- coding: utf-8 -*-

# sudo pip install python-docx
#
# SimpleDocService
#  python-doxに関する簡単なサービスを提供します。
#  まぁ docxライブラリを理解するためのコードですね。
#

from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt

class SimpleDocxService:

    def __init__(self):
        self.document = Document()
        self.latest_run = None

    def set_normal_font(self, name, size):
        # フォントの設定
        font = self.document.styles['Normal'].font
        font.name = name
        font.size = Pt(size)

    def add_head(self, text, lv):
        # 見出しの設定
        self.document.add_heading(text, level=lv)

    def open_text(self):
        # テキスト追加開始
        self.paragraph = self.document.add_paragraph()

    def close_text(self):
        # テキスト追加終了
        return #現状では特に処理はなし

    def get_unicode_text(self, text, src_code):
        # python-docxで扱えるように unicodeに変換
        #return unicode(text, src_code)
        print(type (text))
        return str(text)
        #return text

    def adjust_return_code(self, text):
        # テキストファイルのデータをそのままaddすると改行が
        # 面倒なことになるので、それを削除
        text = text.replace("\n", "")
        text = text.replace("\r", "")
        return text

    def add_text(self, text):
        # テキスト追加
        self.latest_run = self.paragraph.add_run(text)

    def add_text_italic(self, text):
        # テキスト追加（イタリックに）
        self.paragraph.add_run(text).italic = True

    def add_text_bold(self, text):
        # テキスト追加（強調）
        self.paragraph.add_run(text).bold = True

    def add_text_color(self, text, r, g, b):
        # 文字に色をつける
        self.paragraph.add_run(text).font.color.rgb = RGBColor(r, g, b)

    def add_picture(self, filename, inch):
        # 図を挿入する
        self.document.add_picture(filename, width=Inches(inch))

    def save(self, name):
        # docxファイルとして出力。
        self.document.save(name)
